"""
Fills the Upmotion Tech invoice Word template (app/templates/invoice_template.docx)
with a real invoice's data and converts the result to PDF.

The template is a fixed-layout Word document (letterhead, "ISSUED TO"/"INVOICE
NO"/"DATE" fields, and a line-item table are all inside floating textboxes/
tables rather than plain paragraphs) with no merge fields — so instead of a
templating library, this walks every <w:t> text run in the document
(python-docx's element tree reaches into textboxes fine via `.//`) and
replaces known label text in place.

PDF conversion uses docx2pdf, which drives Microsoft Word via COM automation
on Windows (confirmed available in this dev environment). That only works on
a Windows machine with Word installed — it will NOT work on the Linux/Render
production target. Before deploying this, swap `_convert_to_pdf` for a
LibreOffice-headless call (`soffice --headless --convert-to pdf`) or a cloud
conversion API; the docx-filling logic above it is platform-independent and
does not need to change.
"""
import copy
import os
import tempfile
import uuid
from typing import Optional

try:
    import docx
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

from app.models.invoice import Invoice

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "invoice_template.docx")

CURRENCY_SYMBOL = {"USD": "$", "PKR": "₨"}
CURRENCY_WORDS = {"USD": "US Dollars", "PKR": "Pakistani Rupees"}

_ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
         "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
         "Seventeen", "Eighteen", "Nineteen"]
_TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]


def _int_to_words(n: int) -> str:
    if n == 0:
        return "Zero"

    def three_digits(num: int) -> str:
        parts = []
        if num >= 100:
            parts.append(_ONES[num // 100] + " Hundred")
            num %= 100
        if num >= 20:
            parts.append(_TENS[num // 10] + (f"-{_ONES[num % 10]}" if num % 10 else ""))
        elif num > 0:
            parts.append(_ONES[num])
        return " ".join(parts)

    scales = [(1_000_000_000, "Billion"), (1_000_000, "Million"), (1_000, "Thousand")]
    words = []
    remaining = n
    for value, name in scales:
        if remaining >= value:
            words.append(f"{three_digits(remaining // value)} {name}")
            remaining %= value
    if remaining > 0:
        words.append(three_digits(remaining))
    return " ".join(words)


def amount_in_words(amount: float, currency: str) -> str:
    whole = int(amount)
    cents = round((amount - whole) * 100)
    currency_name = CURRENCY_WORDS.get(currency, currency)
    text = f"{_int_to_words(whole)} {currency_name}"
    if cents:
        text += f" and {_int_to_words(cents)} Cents"
    return text + " Only"


def _fmt_money(n: float) -> str:
    return f"{n:,.2f}" if n % 1 else f"{n:,.0f}"


def _set_run_group_text(t_nodes: list, new_text: str, bold: Optional[bool] = None) -> None:
    if not t_nodes:
        return
    t_nodes[0].text = new_text
    for t in t_nodes[1:]:
        t.text = ""
    if bold is not None:
        # t_nodes are <w:t> elements; the run's <w:rPr><w:b/> lives on the
        # parent <w:r>. Force bold on just the first run (the one that now
        # carries all the text) rather than every run in the group.
        run_el = t_nodes[0].getparent()
        rpr = run_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = run_el.makeelement(qn("w:rPr"), {})
            run_el.insert(0, rpr)
        b_el = rpr.find(qn("w:b"))
        if bold:
            if b_el is None:
                rpr.insert(0, rpr.makeelement(qn("w:b"), {}))
        elif b_el is not None:
            rpr.remove(b_el)


def _set_cell_text(cell, text: str, bold: Optional[bool] = None) -> None:
    """Fill a table cell's text WITHOUT going through `Cell.text = ...` —
    that setter replaces all paragraphs with a brand-new default-styled one,
    silently dropping the template's alignment/font/size (confirmed: a
    template cell's CENTER alignment and 10pt run size both come back as
    None after `cell.text = x`). Reusing the first existing run instead
    keeps whatever formatting the template author set up, the same
    preserve-don't-replace approach `_set_run_group_text` already uses for
    the textbox fields above."""
    paragraphs = cell.paragraphs
    target_p = next((p for p in paragraphs if p.runs), paragraphs[0])
    if target_p.runs:
        target_p.runs[0].text = text
        for r in target_p.runs[1:]:
            r.text = ""
        if bold is not None:
            target_p.runs[0].font.bold = bold
    else:
        run = target_p.add_run(text)
        if bold is not None:
            run.font.bold = bold
    for p in paragraphs:
        if p is not target_p:
            for r in p.runs:
                r.text = ""


def _fill_text_fields(body, invoice: Invoice) -> None:
    for t in body.findall(".//" + qn("w:t")):
        stripped = (t.text or "").strip()
        if stripped == "ISSUED TO:":
            t.text = f"ISSUED TO: {invoice.client}"

    for p in body.findall(".//" + qn("w:p")):
        # Direct-child runs only: this template's textboxes have some `<w:p>`
        # elements that structurally contain *other* whole paragraphs as
        # descendants (Word's own compatibility markup) — a deep `.//w:t`
        # search would pick up those nested paragraphs' runs too and corrupt
        # them when a sibling field (e.g. DATE) shares the same paragraph
        # further down. Scoping to `./w:r/w:t` keeps each match to only the
        # runs that actually belong to this exact paragraph.
        t_nodes = p.findall("./" + qn("w:r") + "/" + qn("w:t"))
        if not t_nodes:
            continue
        full_text = "".join(t.text or "" for t in t_nodes).strip()
        if full_text.startswith("INVOICE NO:"):
            _set_run_group_text(t_nodes, f"INVOICE NO: {invoice.invoice_number}")
        elif full_text.startswith("DATE:"):
            date_str = invoice.issue_date.strftime("%d / %m / %Y")
            # No dedicated template slot for a paid date, so it rides on the
            # same line as the issue date — only shown once the invoice is
            # actually marked Paid and a paid date has been recorded.
            if invoice.status == "Paid" and invoice.paid_date:
                date_str += f"     PAID: {invoice.paid_date.strftime('%d / %m / %Y')}"
            _set_run_group_text(t_nodes, f"DATE:  {date_str}")
        elif full_text.startswith("In Words:"):
            total = sum((item.get("qty", 1) or 1) * (item.get("unit_price", 0) or 0) for item in (invoice.line_items or []))
            _set_run_group_text(t_nodes, f"In Words: {amount_in_words(total, invoice.currency)}")
        elif full_text.startswith("NOTES:") and invoice.notes:
            _set_run_group_text(t_nodes, f"NOTES: {invoice.notes}")
        elif full_text.startswith("Account Name:") and invoice.bank_account_name:
            _set_run_group_text(t_nodes, f"Account Name: {invoice.bank_account_name}", bold=True)
        elif full_text.startswith("Account Number:") and invoice.bank_account_number:
            _set_run_group_text(t_nodes, f"Account Number: {invoice.bank_account_number}", bold=True)
        elif full_text.startswith("IBAN Number:") and invoice.bank_iban:
            _set_run_group_text(t_nodes, f"IBAN Number: {invoice.bank_iban}", bold=True)
        elif full_text.startswith("Bank Name:") and invoice.bank_name:
            _set_run_group_text(t_nodes, f"Bank Name: {invoice.bank_name}", bold=True)


def _fill_line_items_table(doc, invoice: Invoice) -> float:
    table = doc.tables[0]
    sym = CURRENCY_SYMBOL.get(invoice.currency, invoice.currency + " ")
    items = invoice.line_items or []
    if not items:
        items = [{"description": invoice.invoice_type or "Services rendered", "qty": 1, "unit_price": invoice.amount}]

    data_rows = table.rows[1:-1]
    total_row_tr = table.rows[-1]._tr
    sample_row_tr = table.rows[1]._tr

    grand_total = 0.0
    for i, item in enumerate(items):
        qty = item.get("qty", 1) or 1
        unit_price = item.get("unit_price", 0) or 0
        line_total = qty * unit_price
        grand_total += line_total

        if i < len(data_rows):
            row = data_rows[i]
        else:
            new_tr = copy.deepcopy(sample_row_tr)
            total_row_tr.addprevious(new_tr)
            row = table.rows[len(table.rows) - 2]

        _set_cell_text(row.cells[0], item.get("description", ""))
        _set_cell_text(row.cells[1], _fmt_money(qty))
        _set_cell_text(row.cells[2], f"{sym}{_fmt_money(unit_price)}")
        _set_cell_text(row.cells[3], f"{sym}{_fmt_money(line_total)}")

    for i in range(len(items), len(data_rows)):
        for c in data_rows[i].cells:
            _set_cell_text(c, "")

    _set_cell_text(table.rows[-1].cells[3], f"{sym} {_fmt_money(grand_total)}")
    return grand_total


def generate_invoice_docx(invoice: Invoice) -> str:
    """Fills the template and saves it to a temp .docx file, returning its path."""
    doc = docx.Document(TEMPLATE_PATH)
    _fill_line_items_table(doc, invoice)
    _fill_text_fields(doc.element.body, invoice)

    tmp_dir = tempfile.gettempdir()
    docx_path = os.path.join(tmp_dir, f"orbit_invoice_{uuid.uuid4().hex}.docx")
    doc.save(docx_path)
    return docx_path


def convert_docx_to_pdf(docx_path: str) -> str:
    """Windows/Word-COM conversion (see module docstring for the production caveat)."""
    from docx2pdf import convert
    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    convert(docx_path, pdf_path)
    return pdf_path


def generate_invoice_pdf_bytes(invoice: Invoice) -> bytes:
    if not _HAS_DOCX:
        raise RuntimeError("PDF generation requires python-docx (install with: pip install python-docx)")
    docx_path = generate_invoice_docx(invoice)
    pdf_path: Optional[str] = None
    try:
        pdf_path = convert_docx_to_pdf(docx_path)
        with open(pdf_path, "rb") as f:
            return f.read()
    finally:
        for path in (docx_path, pdf_path):
            if path and os.path.exists(path):
                os.remove(path)


def safe_invoice_filename(invoice_number: str) -> str:
    cleaned = "".join(c for c in invoice_number if c.isalnum() or c in ("-", "_"))
    return f"{cleaned or 'invoice'}.pdf"
